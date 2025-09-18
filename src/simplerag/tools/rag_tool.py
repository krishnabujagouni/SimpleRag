
import os, hashlib
from typing import List, Dict, Optional, Union
import lancedb
import pandas as pd
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from crewai.tools import tool
import mimetypes
from pathlib import Path
import logging

# Document parsing imports
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

try:
    import easyocr
    import cv2
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


try:
    import openpyxl
   
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    import csv
    CSV_AVAILABLE = True
except ImportError:
    CSV_AVAILABLE = False

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- Config ---
PROVIDER = os.getenv("EMBED_PROVIDER", "gemini")
MODEL = os.getenv("EMBED_MODEL", "text-embedding-004")

# --- Document Parser Class ---
class DocumentParser:
    """Universal document parser that normalizes inputs across formats"""
    
    def __init__(self):
        self.supported_formats = {
            'pdf': self._parse_pdf,
            'docx': self._parse_docx,
            'doc': self._parse_doc,
            'html': self._parse_html,
            'htm': self._parse_html,
            'txt': self._parse_text,
            'md': self._parse_text,
            'csv': self._parse_csv,
            'xlsx': self._parse_excel,
            'xls': self._parse_excel,
            'json': self._parse_json,
            'xml': self._parse_xml,
            'jpg': self._parse_image,
            'jpeg': self._parse_image,
            'png': self._parse_image,
            'bmp': self._parse_image,
            'tiff': self._parse_image,
        }
        
        # Initialize OCR reader if available
        if OCR_AVAILABLE:
            try:
                self.ocr_reader = easyocr.Reader(['en'])
            except Exception as e:
                logger.warning(f"Failed to initialize OCR: {e}")
                self.ocr_reader = None
        else:
            self.ocr_reader = None
    
    def parse_document(self, file_path: str) -> Dict[str, Union[str, List[str], Dict]]:
        """
        Parse a document and return structured data
        
        Returns:
            Dict with keys:
            - 'content': Main text content
            - 'metadata': Document metadata (title, author, etc.)
            - 'structure': Structured data if applicable
            - 'format': Original document format
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Determine file type
            file_extension = file_path.suffix.lower().lstrip('.')
            mime_type, _ = mimetypes.guess_type(str(file_path))
            
            # Check if format is supported
            if file_extension not in self.supported_formats:
                logger.warning(f"Unsupported format: {file_extension}, treating as text")
                file_extension = 'txt'
            
            # Parse the document
            parser_func = self.supported_formats[file_extension]
            result = parser_func(file_path)
            
            # Add metadata
            result['format'] = file_extension
            result['file_path'] = str(file_path)
            result['file_size'] = file_path.stat().st_size
            result['mime_type'] = mime_type
            
            return result
            
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {e}")
            return {
                'content': '',
                'metadata': {'error': str(e)},
                'structure': {},
                'format': 'error'
            }
    
    def _parse_pdf(self, file_path: Path) -> Dict:
        """Parse PDF documents"""
        if not PDF_AVAILABLE:
            raise ImportError("PDF parsing requires PyPDF2 and pdfplumber: pip install PyPDF2 pdfplumber")
        
        content = ""
        metadata = {}
        structure = {'pages': []}
        
        try:
            # Try pdfplumber first (better text extraction)
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                metadata.update({
                    'pages': len(pdf.pages),
                    'title': getattr(pdf.metadata, 'title', ''),
                    'author': getattr(pdf.metadata, 'author', ''),
                    'creator': getattr(pdf.metadata, 'creator', '')
                })
                
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    content += page_text + "\n\n"
                    structure['pages'].append({
                        'page_number': i + 1,
                        'text': page_text,
                        'tables': page.extract_tables() or []
                    })
                    
        except Exception as e:
            # Fallback to PyPDF2
            logger.warning(f"pdfplumber failed, using PyPDF2: {e}")
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata.update({
                    'pages': len(pdf_reader.pages),
                    'title': getattr(pdf_reader.metadata, '/Title', ''),
                    'author': getattr(pdf_reader.metadata, '/Author', ''),
                    'creator': getattr(pdf_reader.metadata, '/Creator', '')
                })
                
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    content += page_text + "\n\n"
                    structure['pages'].append({
                        'page_number': i + 1,
                        'text': page_text
                    })
        
        return {
            'content': content.strip(),
            'metadata': metadata,
            'structure': structure
        }
    
    def _parse_docx(self, file_path: Path) -> Dict:
        """Parse DOCX documents"""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX parsing requires python-docx: pip install python-docx")
        
        doc = DocxDocument(file_path)
        
        # Extract text content
        content = ""
        paragraphs = []
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                content += para_text + "\n"
                paragraphs.append({
                    'text': para_text,
                    'style': para.style.name if para.style else None
                })
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        # Extract metadata
        metadata = {
            'title': doc.core_properties.title or '',
            'author': doc.core_properties.author or '',
            'subject': doc.core_properties.subject or '',
            'created': str(doc.core_properties.created) if doc.core_properties.created else '',
            'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
            'paragraphs_count': len(paragraphs),
            'tables_count': len(tables)
        }
        
        return {
            'content': content.strip(),
            'metadata': metadata,
            'structure': {
                'paragraphs': paragraphs,
                'tables': tables
            }
        }
    
    def _parse_doc(self, file_path: Path) -> Dict:
        """Parse legacy DOC documents (requires additional tools)"""
        # For .doc files, you might need antiword or libreoffice
        # This is a placeholder - implement based on your needs
        try:
            import subprocess
            result = subprocess.run(['antiword', str(file_path)], 
                                  capture_output=True, text=True, check=True)
            content = result.stdout
        except (subprocess.SubprocessError, FileNotFoundError):
            logger.warning("antiword not available, treating .doc as text")
            content = self._parse_text(file_path)['content']
        
        return {
            'content': content,
            'metadata': {'format': 'doc'},
            'structure': {}
        }
    
    def _parse_html(self, file_path: Path) -> Dict:
        """Parse HTML documents"""
        if not HTML_AVAILABLE:
            raise ImportError("HTML parsing requires BeautifulSoup: pip install beautifulsoup4")
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            html_content = file.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extract metadata
        metadata = {
            'title': soup.title.string if soup.title else '',
            'description': '',
            'keywords': ''
        }
        
        # Get meta tags
        for meta in soup.find_all('meta'):
            if meta.get('name') == 'description':
                metadata['description'] = meta.get('content', '')
            elif meta.get('name') == 'keywords':
                metadata['keywords'] = meta.get('content', '')
        
        # Extract structured content
        structure = {
            'headings': [],
            'links': [],
            'images': []
        }
        
        # Get headings
        for i in range(1, 7):
            for heading in soup.find_all(f'h{i}'):
                structure['headings'].append({
                    'level': i,
                    'text': heading.get_text().strip()
                })
        
        # Get links
        for link in soup.find_all('a', href=True):
            structure['links'].append({
                'text': link.get_text().strip(),
                'href': link['href']
            })
        
        # Get images
        for img in soup.find_all('img'):
            structure['images'].append({
                'src': img.get('src', ''),
                'alt': img.get('alt', '')
            })
        
        # Extract clean text
        for script in soup(["script", "style"]):
            script.decompose()
        content = soup.get_text()
        
        # Clean up text
        lines = (line.strip() for line in content.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        content = ' '.join(chunk for chunk in chunks if chunk)
        
        return {
            'content': content,
            'metadata': metadata,
            'structure': structure
        }
    
    def _parse_text(self, file_path: Path) -> Dict:
        """Parse plain text documents"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            content = file.read()
        
        return {
            'content': content,
            'metadata': {'lines': len(content.splitlines())},
            'structure': {}
        }
    
    def _parse_csv(self, file_path: Path) -> Dict:
        """Parse CSV files"""
        if not CSV_AVAILABLE:
            import csv as csv_module
        
        content = ""
        structure = {'headers': [], 'rows': []}
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            csv_reader = csv.reader(file)
            rows = list(csv_reader)
            
            if rows:
                structure['headers'] = rows[0]
                structure['rows'] = rows[1:] if len(rows) > 1 else []
                
                # Convert to readable text
                content = f"Headers: {', '.join(structure['headers'])}\n\n"
                for i, row in enumerate(structure['rows'][:100]):  # Limit to first 100 rows
                    content += f"Row {i+1}: {', '.join(row)}\n"
        
        return {
            'content': content,
            'metadata': {
                'rows': len(structure['rows']),
                'columns': len(structure['headers'])
            },
            'structure': structure
        }
    
    def _parse_excel(self, file_path: Path) -> Dict:
        """Parse Excel files"""
        if not EXCEL_AVAILABLE:
            raise ImportError("Excel parsing requires openpyxl: pip install openpyxl")
        
        import openpyxl
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        
        content = ""
        structure = {'sheets': []}
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_data = []
            
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    sheet_data.append([str(cell) if cell is not None else '' for cell in row])
            
            structure['sheets'].append({
                'name': sheet_name,
                'data': sheet_data
            })
            
            # Add to content
            content += f"Sheet: {sheet_name}\n"
            for i, row in enumerate(sheet_data[:50]):  # Limit rows
                content += f"Row {i+1}: {', '.join(row)}\n"
            content += "\n"
        
        return {
            'content': content,
            'metadata': {'sheets': len(workbook.sheetnames)},
            'structure': structure
        }
    
    def _parse_json(self, file_path: Path) -> Dict:
        """Parse JSON files"""
        import json
        
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        
        # Convert JSON to readable text
        content = json.dumps(data, indent=2, ensure_ascii=False)
        
        return {
            'content': content,
            'metadata': {'type': type(data).__name__},
            'structure': {'json_data': data}
        }
    
    def _parse_xml(self, file_path: Path) -> Dict:
        """Parse XML files"""
        if not HTML_AVAILABLE:
            raise ImportError("XML parsing requires BeautifulSoup: pip install beautifulsoup4")
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            xml_content = file.read()
        
        soup = BeautifulSoup(xml_content, 'xml')
        content = soup.get_text()
        
        return {
            'content': content,
            'metadata': {'root_tag': soup.name if soup.name else ''},
            'structure': {}
        }
    
    def _parse_image(self, file_path: Path) -> Dict:
        """Parse images using OCR"""
        if not OCR_AVAILABLE or not self.ocr_reader:
            return {
                'content': f"[IMAGE: {file_path.name}]",
                'metadata': {'type': 'image', 'ocr_available': False},
                'structure': {}
            }
        
        try:
            result = self.ocr_reader.readtext(str(file_path))
            content = "\n".join([text[1] for text in result if text[2] > 0.5])  # Confidence > 0.5
            
            return {
                'content': content,
                'metadata': {
                    'type': 'image',
                    'ocr_results': len(result),
                    'avg_confidence': sum(text[2] for text in result) / len(result) if result else 0
                },
                'structure': {'ocr_data': result}
            }
        except Exception as e:
            logger.error(f"OCR failed for {file_path}: {e}")
            return {
                'content': f"[IMAGE: {file_path.name}]",
                'metadata': {'type': 'image', 'ocr_error': str(e)},
                'structure': {}
            }
    
    # def _parse_audio(self, file_path: Path) -> Dict:
    #     """Parse audio files using speech recognition"""
    #     if not AUDIO_AVAILABLE:
    #         return {
    #             'content': f"[AUDIO: {file_path.name}]",
    #             'metadata': {'type': 'audio', 'transcription_available': False},
    #             'structure': {}
    #         }
        
    #     try:
    #         # Convert audio to wav if needed
    #         audio = pydub.AudioSegment.from_file(str(file_path))
            
    #         # Use speech recognition
    #         r = sr.Recognizer()
    #         with sr.AudioFile(str(file_path)) as source:
    #             audio_data = r.record(source)
    #             content = r.recognize_google(audio_data)
            
    #         return {
    #             'content': content,
    #             'metadata': {
    #                 'type': 'audio',
    #                 'duration': len(audio) / 1000.0,  # seconds
    #                 'sample_rate': audio.frame_rate
    #             },
    #             'structure': {}
    #         }
    #     except Exception as e:
    #         logger.error(f"Audio transcription failed for {file_path}: {e}")
    #         return {
    #             'content': f"[AUDIO: {file_path.name}]",
    #             'metadata': {'type': 'audio', 'transcription_error': str(e)},
    #             'structure': {}
    #         }

# --- Hash ID helper
def _make_id(text: str) -> str:
    return hashlib.md5(text.encode("utf-8")).hexdigest()

# --- Universal Embedding Wrapper
def embed_text(text: str) -> List[float]:
    """Embed text with user-selected provider."""
    if PROVIDER == "gemini":
        import google.generativeai as genai
        try:
            api_key = os.getenv("GOOGLE_API_KEY")
            if not api_key:
                raise ValueError("GOOGLE_API_KEY environment variable is required for Gemini")
            
            genai.configure(api_key=api_key)
            
            if MODEL == "text-embedding-004":
                resp = genai.embed_content(
                    model=f"models/{MODEL}",
                    content=text,
                    task_type="retrieval_document"
                )
            else:
                resp = genai.embed_content(
                    model=MODEL if MODEL.startswith("models/") else f"models/{MODEL}",
                    content=text,
                    task_type="retrieval_document"
                )
            return resp["embedding"]
        except Exception as e:
            logger.error(f"Error calling Gemini API: {e}")
            raise
    else:
        raise ValueError(f"Unsupported provider: {PROVIDER}")

def embed_query(query: str) -> List[float]:
    """Embed query with user-selected provider."""
    if PROVIDER == "gemini":
        import google.generativeai as genai
        try:
            api_key = os.getenv("GOOGLE_API_KEY")
            genai.configure(api_key=api_key)
            
            if MODEL == "text-embedding-004":
                resp = genai.embed_content(
                    model=f"models/{MODEL}",
                    content=query,
                    task_type="retrieval_query"
                )
            else:
                resp = genai.embed_content(
                    model=MODEL if MODEL.startswith("models/") else f"models/{MODEL}",
                    content=query,
                    task_type="retrieval_query"
                )
            return resp["embedding"]
        except Exception as e:
            logger.error(f"Error calling Gemini API: {e}")
            raise
    else:
        return embed_text(query)

# --- Enhanced chunking with metadata preservation
def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50, preserve_structure: bool = True) -> List[Dict]:
    """Enhanced chunking that preserves document structure"""
    if not preserve_structure:
        # Simple chunking
        chunks = []
        for i in range(0, len(text), chunk_size - overlap):
            chunk_text = text[i:i + chunk_size]
            chunks.append({
                'text': chunk_text,
                'start_pos': i,
                'end_pos': min(i + chunk_size, len(text)),
                'chunk_id': i // (chunk_size - overlap)
            })
        return chunks
    
    # Structure-aware chunking (split by paragraphs first)
    paragraphs = text.split('\n\n')
    chunks = []
    current_chunk = ""
    chunk_id = 0
    start_pos = 0
    
    for para in paragraphs:
        if len(current_chunk) + len(para) + 2 <= chunk_size:
            current_chunk += para + "\n\n"
        else:
            if current_chunk:
                chunks.append({
                    'text': current_chunk.strip(),
                    'start_pos': start_pos,
                    'end_pos': start_pos + len(current_chunk),
                    'chunk_id': chunk_id
                })
                chunk_id += 1
                start_pos += len(current_chunk)
            current_chunk = para + "\n\n"
    
    # Add final chunk
    if current_chunk:
        chunks.append({
            'text': current_chunk.strip(),
            'start_pos': start_pos,
            'end_pos': start_pos + len(current_chunk),
            'chunk_id': chunk_id
        })
    
    return chunks

# --- Enhanced Core Functions
def _rag_ingest_text_function(text: str, chunk_size: int = 500) -> str:
    """Simple text ingestion (like original code) - no file parsing"""
    try:
        # Simple chunking like the original code
        chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
        embeddings = [embed_text(c) for c in chunks]
        
        # Create DataFrame with simple structure (like original)
        data = []
        for chunk, emb in zip(chunks, embeddings):
            data.append({
                "id": _make_id(chunk),
                "text": chunk,
                "embedding": emb
            })
        
        df = pd.DataFrame(data)
        
        # Connect to LanceDB
        db = lancedb.connect(os.getenv("VECTOR_DB_PATH", "./.lancedb"))
        
        # Create or append to table
        if "documents" in db.table_names():
            tbl = db.open_table("documents")
            tbl.add(df)
        else:
            tbl = db.create_table("documents", df)
        
        return f"✅ Ingested {len(data)} chunks with {len(embeddings[0])}-dim embeddings using {PROVIDER}/{MODEL}."
        
    except Exception as e:
        return f"❌ Error ingesting text: {str(e)}"

def _rag_ingest_file_function(file_path: str, chunk_size: int = 500, overlap: int = 50) -> str:
    """Enhanced file ingest with multi-format support - Compatible with existing schema"""
    try:
        parser = DocumentParser()
        parsed_doc = parser.parse_document(file_path)
        
        if not parsed_doc['content']:
            return f"❌ No content extracted from {file_path}"
        
        # Enhanced chunking with metadata
        chunks_data = chunk_text(
            parsed_doc['content'], 
            chunk_size=chunk_size, 
            overlap=overlap,
            preserve_structure=True
        )
        
        # Generate embeddings for all chunks
        embeddings = []
        for chunk_data in chunks_data:
            emb = embed_text(chunk_data['text'])
            embeddings.append(emb)
        
        # Create DataFrame with SIMPLE schema (compatible with existing table)
        data = []
        for chunk_data, emb in zip(chunks_data, embeddings):
            # Add file info to the text content for reference
            enhanced_text = f"[File: {Path(file_path).name}]\n{chunk_data['text']}"
            data.append({
                "id": _make_id(chunk_data['text']),
                "text": enhanced_text,  # Include filename in text
                "embedding": emb
            })
        
        df = pd.DataFrame(data)
        
        # Connect to LanceDB
        db = lancedb.connect(os.getenv("VECTOR_DB_PATH", "./.lancedb"))
        
        # Create or append to table
        if "documents" in db.table_names():
            tbl = db.open_table("documents")
            tbl.add(df)
        else:
            tbl = db.create_table("documents", df)
        
        return f"✅ Successfully ingested {len(data)} chunks from {parsed_doc['format']} file ({Path(file_path).name}) with {len(embeddings[0])}-dim embeddings using {PROVIDER}/{MODEL}."
        
    except Exception as e:
        logger.error(f"Error ingesting document: {e}")
        return f"❌ Error ingesting document: {str(e)}"

def _rag_retrieve_function(query: str, top_k: int = 3, file_format_filter: Optional[str] = None) -> List[Dict]:
    """Enhanced retrieve with format filtering"""
    try:
        db = lancedb.connect(os.getenv("VECTOR_DB_PATH", "./.lancedb"))
        
        if "documents" not in db.table_names():
            return [{"error": "No documents have been ingested yet. Please ingest some documents first."}]
        
        tbl = db.open_table("documents")
        df = tbl.to_pandas()
        
        if len(df) == 0:
            return [{"error": "No documents found in the database."}]
        
        # Apply format filter if specified
        # if file_format_filter:
        #     df = df[df['file_format'] == file_format_filter]
        #     if len(df) == 0:
        #         return [{"error": f"No documents found with format: {file_format_filter}"}]
        
        # logger.info(f"Found {len(df)} documents in database")
        
        query_embedding = embed_query(query)
        
        try:
            # Try LanceDB vector search
            search_results = tbl.search(query_embedding, vector_column_name="embedding").limit(top_k)
            if file_format_filter:
                search_results = search_results.where(f"file_format = '{file_format_filter}'")
            
            results = search_results.to_list()
            
            formatted_results = []
            for i, result in enumerate(results):
                formatted_results.append({
                    "rank": i + 1,
                    "text": result.get("text", ""),
                    "source_file": result.get("source_file", ""),
                    "file_format": result.get("file_format", ""),
                    "chunk_id": result.get("chunk_id", ""),
                    "distance": result.get("_distance", 0.0),
                    "metadata": result.get("metadata", "{}")
                })
            return formatted_results
            
        except Exception as lance_error:
            logger.warning(f"LanceDB search failed: {lance_error}, using fallback")
            
            # Fallback to manual similarity
            doc_embeddings = df['embedding'].tolist()
            query_vec = np.array(query_embedding).reshape(1, -1)
            doc_vecs = np.array(doc_embeddings)
            
            similarities = cosine_similarity(query_vec, doc_vecs)[0]
            top_indices = np.argsort(similarities)[::-1][:top_k]
            
            formatted_results = []
            for i, idx in enumerate(top_indices):
                row = df.iloc[idx]
                formatted_results.append({
                    "rank": i + 1,
                    "text": row['text'],
                    "source_file": row['source_file'],
                    "file_format": row['file_format'],
                    "chunk_id": row['chunk_id'],
                    "similarity": float(similarities[idx]),
                    "metadata": row['metadata']
                })
            
            return formatted_results
        
    except Exception as e:
        logger.error(f"Error retrieving documents: {e}")
        return [{"error": f"Error retrieving documents: {str(e)}"}]

# --- Tool Decorated Functions (for CrewAI)
@tool
def rag_ingest(document: str) -> str:
    """Split, embed, and store document text in LanceDB (simple text ingestion like original)."""
    return _rag_ingest_text_function(document)

@tool
def rag_ingest_file(file_path: str, chunk_size: int = 500, overlap: int = 50) -> str:
    """Ingest and parse documents from multiple formats (PDF, DOCX, HTML, images, audio, etc.)."""
    return _rag_ingest_file_function(file_path, chunk_size, overlap)

@tool
def rag_retrieve(query: str, top_k: int = 3, file_format: Optional[str] = None) -> List[Dict]:
    """Retrieve relevant chunks with optional format filtering."""
    return _rag_retrieve_function(query, top_k, file_format)

# --- Route Helper Functions
def ingest_text(text: str, api_key: str) -> str:
    """
    Direct text ingestion (exactly like original code)
    Calls the simple text ingest function, not the file parser.
    """
    try:
        return _rag_ingest_text_function(text)
    except Exception as e:
        return f"❌ Failed to ingest text: {e}"

def ingest(file_path: str, api_key: str) -> str:
    """
    File-based ingestion with multi-format parsing
    Uses enhanced parser for various file formats.
    """
    return _rag_ingest_file_function(file_path)

def retrieve(query: str, api_key: str, file_format: Optional[str] = None) -> List[Dict]:
    """Enhanced retrieval with format filtering"""
    return _rag_retrieve_function(query, file_format_filter=file_format)

# --- Utility function to check available parsers
def get_supported_formats() -> Dict[str, bool]:
    """Return dictionary of supported formats and their availability"""
    return {
        'PDF': PDF_AVAILABLE,
        'DOCX': DOCX_AVAILABLE,}